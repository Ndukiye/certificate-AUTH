import csv
from mailmerge import MailMerge
import os
import sys
import argparse
import zipfile
import shutil
import tempfile
import xml.etree.ElementTree as ET
import re
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from docx.shared import Inches
from lxml import etree
import os
import argparse
from docx.oxml import OxmlElement
from docx.text.run import Run

## Removed auto-update-on-open to prevent DOCX corruption warnings

def _embed_qr_codes_in_docx(docx_path: str, paths: list):
    if not paths:
        return

    document = Document(docx_path)
    qr_path_index = 0

    # Namespace map for easier XPath queries
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    }

    # Create XPath objects with namespaces
    fld_char_begin_xpath = etree.XPath('.//w:fldChar[@w:fldCharType="begin"]', namespaces=nsmap)
    fld_char_end_xpath = etree.XPath('.//w:fldChar[@w:fldCharType="end"]', namespaces=nsmap)
    instr_text_xpath = etree.XPath('.//w:instrText', namespaces=nsmap)
    drawing_xpath = etree.XPath('.//w:drawing', namespaces=nsmap)

    for paragraph in document.paragraphs:
        p_element = paragraph._element
        
        # Find all field codes in the paragraph
        field_starts = fld_char_begin_xpath(p_element)
        field_ends = fld_char_end_xpath(p_element)

        # Iterate through field starts and their corresponding ends
        for i, fld_start in enumerate(field_starts):
            if i >= len(field_ends):
                print("Warning: Mismatched field start and end characters in a paragraph.")
                break

            fld_end = field_ends[i]
            
            # Collect all instrText elements between fld_start and fld_end
            full_instr_text = []
            collecting = False
            
            # Iterate through all descendants of the paragraph element in document order
            for element in p_element.iter():
                if element == fld_start:
                    collecting = True
                    continue # Don't add fld_start itself
                
                if element == fld_end:
                    collecting = False
                    break # Stop collecting after fld_end
                
                if collecting and element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instrText':
                    if element.text:
                        full_instr_text.append(element.text)

            field_instruction = "".join(full_instr_text)

            is_qr_field = False
            if "INCLUDEPICTURE" in field_instruction and "MERGEFIELD QRCodePath" in field_instruction:
                is_qr_field = True

            if is_qr_field:
                if qr_path_index < len(paths):
                    image_path = paths[qr_path_index]
                    if os.path.exists(image_path):
                        # Find the parent <w:p> element
                        parent_p = fld_start.getparent()
                        while parent_p is not None and not parent_p.tag.endswith('p'):
                            parent_p = parent_p.getparent()

                        if parent_p is not None:
                            # Find the actual run elements containing fld_start and fld_end
                            start_run = fld_start.getparent()
                            while start_run is not None and not start_run.tag.endswith('r'):
                                start_run = start_run.getparent()

                            end_run = fld_end.getparent()
                            while end_run is not None and not end_run.tag.endswith('r'):
                                end_run = end_run.getparent()

                            if start_run is not None and end_run is not None:
                                all_runs_in_paragraph = list(parent_p.iterchildren())
                                start_run_index = -1
                                end_run_index = -1

                                for idx, run_element in enumerate(all_runs_in_paragraph):
                                    if run_element == start_run:
                                        start_run_index = idx
                                    if run_element == end_run:
                                        end_run_index = idx
                                        break

                                if start_run_index != -1 and end_run_index != -1:
                                    # Remove runs in reverse order to avoid index issues
                                    for i in range(end_run_index, start_run_index - 1, -1):
                                        parent_p.remove(all_runs_in_paragraph[i])

                                    # Create a new run element for the image
                                    new_run_element = OxmlElement('w:r')
                                    parent_p.insert(start_run_index, new_run_element) # Insert at the original start position

                                    # Create a Run object from this new element and add the picture
                                    new_run = Run(new_run_element, paragraph)
                                    new_run.add_picture(image_path, width=Inches(1), height=Inches(1))
                                    qr_path_index += 1
                                else:
                                    print(f"Warning: Could not find start/end run indices for field: {field_instruction}")
                            else:
                                print(f"Warning: Could not find start/end run elements for field: {field_instruction}")
                        else:
                            print(f"Warning: Could not find parent paragraph for field: {field_instruction}")
                    else:
                        print(f"Warning: QR code image not found at {image_path} for field: {field_instruction}")
                else:
                    print(f"Warning: Ran out of QR code paths to embed for field: {field_instruction}")
            else:
                # If it's not a QR field, but it's a field, remove it to prevent errors
                parent_p = fld_start.getparent()
                while parent_p is not None and not parent_p.tag.endswith('p'):
                    parent_p = parent_p.getparent()

                if parent_p is not None:
                    # Find the actual run elements containing fld_start and fld_end
                    start_run = fld_start.getparent()
                    while start_run is not None and not start_run.tag.endswith('r'):
                        start_run = start_run.getparent()

                    end_run = fld_end.getparent()
                    while end_run is not None and not end_run.tag.endswith('r'):
                        end_run = end_run.getparent()

                    if start_run is not None and end_run is not None:
                        all_runs_in_paragraph = list(parent_p.iterchildren())
                        start_run_index = -1
                        end_run_index = -1

                        for idx, run_element in enumerate(all_runs_in_paragraph):
                            if run_element == start_run:
                                start_run_index = idx
                            if run_element == end_run:
                                end_run_index = idx
                                break

                        if start_run_index != -1 and end_run_index != -1:
                            # Remove runs in reverse order to avoid index issues
                            for i in range(end_run_index, start_run_index - 1, -1):
                                parent_p.remove(all_runs_in_paragraph[i])
                        else:
                            print(f"Warning: Could not find start/end run indices for non-QR field: {field_instruction}")
                    else:
                        print(f"Warning: Could not find start/end run elements for non-QR field: {field_instruction}")
                else:
                    print(f"Warning: Could not find parent paragraph for non-QR field: {field_instruction}")
    # Fallback: if not all QR codes were embedded via field codes, insert by marker text
    try:
        for paragraph in document.paragraphs:
            if qr_path_index >= len(paths):
                break
            # Use a reliable marker text already present in the template
            if 'Scan QR Code to Verify' in paragraph.text:
                # Skip if an image is already in this paragraph
                has_image = len(drawing_xpath(paragraph._element)) > 0
                if not has_image:
                    image_path = paths[qr_path_index]
                    if os.path.exists(image_path):
                        run_el = OxmlElement('w:r')
                        paragraph._element.append(run_el)
                        run = Run(run_el, paragraph)
                        run.add_picture(image_path, width=Inches(1), height=Inches(1))
                        qr_path_index += 1
    except Exception as e:
        print(f"Warning: Fallback QR insertion by marker failed: {e}")

    # Cleanup: remove any lingering field-code runs globally (fldChar/instrText)
    try:
        _remove_field_code_runs(document)
    except Exception as e:
        print(f"Warning: Cleanup of field codes failed: {e}")

    document.save(docx_path)
    print("Successfully embedded QR code images.")

def _inline_includepicture_paths(docx_path: str, paths: list):
    # This function is now deprecated as _embed_qr_codes_in_docx handles the embedding
    # Keeping it for now, but it will not be called.
    print("Warning: _inline_includepicture_paths is deprecated and will not be executed.")
    return

def _remove_field_code_runs(document: Document):
    # Remove any w:r that contains w:fldChar or w:instrText anywhere in the document
    for paragraph in document.paragraphs:
        p = paragraph._element
        # Collect runs first to avoid concurrent modification issues
        runs = list(p.iterchildren())
        for r in runs:
            try:
                # Scan descendants for fldChar or instrText
                remove = False
                for el in r.iter():
                    tag = el.tag if isinstance(el.tag, str) else ''
                    if tag.endswith('fldChar') or tag.endswith('instrText'):
                        remove = True
                        break
                if remove:
                    p.remove(r)
            except Exception:
                # If anything odd occurs, skip this run
                continue

def main(input_csv_path=None, template_path=None, output_docx_path=None, base_path=None):
    print(f"export_merge_docx.py main called with: input_csv_path={input_csv_path}, template_path={template_path}, output_docx_path={output_docx_path}, base_path={base_path}")

    if base_path is None:
        base_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

    if input_csv_path is None:
        input_csv_path = os.path.join(base_path, 'data', 'Certificates_Chained.csv')
    if template_path is None:
        template_path = os.path.join(base_path, 'templates', 'certificate_template.docx')
    if output_docx_path is None:
        output_docx_path = os.path.join(base_path, 'data', 'Certificates_Merged.docx')

    # Ensure paths are absolute and correctly resolved
    # input_csv_path = os.path.abspath(input_csv_path)
    # template_path = os.path.abspath(template_path)
    # output_docx_path = os.path.abspath(output_docx_path)
    # base_path = os.path.abspath(base_path)

    print(f"Resolved paths: input_csv_path={input_csv_path}, template_path={template_path}, output_docx_path={output_docx_path}, base_path={base_path}")

    if not os.path.exists(input_csv_path):
        msg = f"Input CSV file not found at {input_csv_path}"
        print(f"Error: {msg}")
        raise FileNotFoundError(msg)
    if not os.path.exists(template_path):
        msg = f"Template DOCX file not found at {template_path}"
        print(f"Error: {msg}")
        raise FileNotFoundError(msg)

    try:
        with open(input_csv_path, 'r', encoding='utf-8') as csv_file:
            csv_reader = csv.DictReader(csv_file)
            records = list(csv_reader)

        if not records:
            print("No records found in the CSV file. Writing original template to output and exiting.")
            # Write the original template to the output path to avoid frontend download issues
            document = MailMerge(template_path)
            # Ensure the output directory exists
            output_dir = os.path.dirname(output_docx_path)
            os.makedirs(output_dir, exist_ok=True)
            document.write(output_docx_path)
            print(f"No data to merge. Output saved to {output_docx_path}")
            return

        # Assuming all records use the same template
        document = MailMerge(template_path)
        
        # Get all merge fields in the document
        merge_fields = document.get_merge_fields()
        print(f"Merge fields found in template: {merge_fields}")

        # Debugging: Print the raw merge fields from the document
        print(f"Raw merge fields from document: {document.get_merge_fields()}")

        # Prepare a list of dictionaries for merging
        merge_data = []
        for record in records:
            # Create a dictionary for each record, ensuring all merge fields are present
            record_data = {}
            for field in merge_fields:
                # Map template merge fields to CSV column headers
                csv_column_map = {
                    'ID': 'CertID',
                    'CompletionDate': 'DateIssued',
                    'Hash': 'CurrentHash',
                    'Link': 'VerificationURL',
                    'QR': 'QRCodeURL',
                    'RecipientName': 'RecipientName',
                    'CourseTitle': 'CourseTitle'
                }
                csv_column_name = csv_column_map.get(field, field) # Use mapping, or field name if not mapped
                
                if field == 'QR':
                    qr_relative_path = record.get('QRCodePath', '')
                    if qr_relative_path:
                        # If QRCodePath is absolute, join will return the absolute path; otherwise resolve relative to base_path
                        record_data[field] = os.path.join(base_path, qr_relative_path)
                    else:
                        record_data[field] = ''
                else:
                    record_data[field] = record.get(csv_column_name, '') # Use empty string if column not in CSV
            
            # Also set QRCodePath explicitly for templates that use INCLUDEPICTURE with MERGEFIELD QRCodePath
            qr_path_value = record.get('QRCodePath', '')
            if qr_path_value:
                record_data['QRCodePath'] = os.path.join(base_path, qr_path_value)
            else:
                record_data['QRCodePath'] = ''

            # Ensure VerificationURL is available even if template uses a different field name
            if 'VerificationURL' not in record_data:
                record_data['VerificationURL'] = record.get('VerificationURL', '')

            # Explicitly add QR field if not already present (e.g., if get_merge_fields() doesn't find it)
            if 'QR' not in merge_fields:
                qr_relative_path = record.get('QRCodePath', '')
                if qr_relative_path:
                    record_data['QR'] = os.path.join(base_path, qr_relative_path)
                else:
                    record_data['QR'] = ''
            merge_data.append(record_data)

        # Merge all records
        document.merge_templates(merge_data, separator='page_break')

        # Ensure the output directory exists
        output_dir = os.path.dirname(output_docx_path)
        os.makedirs(output_dir, exist_ok=True)

        document.write(output_docx_path)
        
        # Post-process: replace nested MERGEFIELD QRCodePath in INCLUDEPICTURE with literal paths
        try:
            qr_paths = []
            for r in records:
                p = r.get('QRCodePath', '')
                if p:
                    qr_paths.append(os.path.join(base_path, p))
            _embed_qr_codes_in_docx(output_docx_path, qr_paths)
            print("Successfully embedded QR code images into the document.")
        except Exception as e:
            print(f"Warning: Failed to embed QR code images: {e}")

        # Note: Word fields may require manual refresh (Ctrl+A, F9) after merge
        print(f"Mail merge completed successfully. Output saved to {output_docx_path}")

    except Exception as e:
        import traceback
        print("An error occurred during mail merge:\n" + traceback.format_exc())
        # Re-raise so the Flask route can handle and return a proper 500 JSON
        raise


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Perform mail merge from CSV to DOCX template.')
    parser.add_argument('--input-csv', type=str, help='Path to the input CSV file.')
    parser.add_argument('--template', type=str, help='Path to the DOCX template file.')
    parser.add_argument('--output-docx', type=str, help='Path to the output DOCX file.')
    parser.add_argument('--base-path', type=str, help='Base path for resolving relative paths, e.g., for QR codes.')
    args = parser.parse_args()

    main(args.input_csv, args.template, args.output_docx, args.base_path)